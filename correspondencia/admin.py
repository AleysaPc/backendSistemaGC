from datetime import timezone
from django.contrib import admin
from django.http import HttpResponse
from docx import Document
from io import BytesIO
from .models import Correspondencia, CorrespondenciaEntrante, CorrespondenciaSaliente, TipoDocumento
from .models import FlujoAprobacion, DescargaDocumento


# Configuración personalizada para CorrespondenciaEntrante
@admin.register(CorrespondenciaEntrante)
class CorrespondenciaEntranteAdmin(admin.ModelAdmin):
    # Excluye el campo nro_registro del formulario
    exclude = ('nro_registro',)

    # Si deseas mostrar el campo como de solo lectura, agrégalo a readonly_fields
    readonly_fields = ('nro_registro',)

    # Campos que se mostrarán en la lista de registros en el admin
    list_display = ('nro_registro', 'fecha_recepcion_formateada', 'fecha_respuesta_formateada', 'mostrar_referencia', 'mostrar_remitente')

    # Metodo para mostrar la referencia
    def mostrar_referencia(self, obj):
        # Acceder al campo 'referencia' del modelo relacionado 'Correspondencia'
        return obj.referencia
    mostrar_referencia.short_description = 'Referencia'

    #Metodo para mostrar remitente
    def mostrar_remitente(self, obj):
       remitente = obj.remitente
       institucion = remitente.institucion 
       return f"{remitente.nombre} - {remitente.apellido} - {institucion.nombre}"
    mostrar_remitente.short_description = 'Remitente'

    def fecha_recepcion_formateada(self, obj):
        if obj.fecha_recepcion:
            return obj.fecha_recepcion.strftime("%d-%m-%Y %H:%M")  # Formato: "01-03-2025 14:30"
        return ""
    fecha_recepcion_formateada.short_description = 'Fecha de Recepción'

    def fecha_respuesta_formateada(self, obj):
        if obj.fecha_respuesta:
            return obj.fecha_respuesta.strftime("%d-%m-%Y %H:%M")  # Formato: "01-03-2025 14:30"
        return ""
    fecha_respuesta_formateada.short_description = 'Fecha de Respuesta'

@admin.register(CorrespondenciaSaliente)
class CorrespondenciaSalienteAdmin(admin.ModelAdmin):
    list_display = ('cite', 'referencia', 'remitente', 'fecha_envio', 'estado')
    fields = ('cite', 'fecha_envio', 'remitente', 'referencia', 'descripcion', 
              'prioridad', 'estado', 'personal_destinatario', 'archivo_word')
    readonly_fields = ('cite',)

    def formfield_for_dbfield(self, db_field, request, **kwargs):
        if db_field.name == "remitente":
            kwargs["label"] = "Destinatario"  # Cambiar el nombre en la interfaz
        return super().formfield_for_dbfield(db_field, request, **kwargs)

    def get_fields(self, request, obj=None):
        base_fields = super().get_fields(request, obj)

        # Si la correspondencia está aprobada, se agregan campos adicionales y se excluye 'archivo_word'
        if obj and obj.estado == 'aprobado':
            return tuple(field for field in base_fields if field != 'archivo_word') + (
                'fecha_recepcion', 'fecha_seguimiento', 'paginas', 'documento'
            )

        return base_fields

           
    actions = ['generar_documento_word',]
    def generar_documento_word(self, request, queryset):
        if queryset.count() != 1:
            self.message_user(request, "Solo puedes generar un documento a la vez.", level='error')
            return

        correspondencia_saliente = queryset.first()

        # Crear el archivo Word
        doc = Document()
        doc.add_paragraph(f"La Paz {correspondencia_saliente.fecha_envio.strftime('%Y-%m-%d')}")
        doc.add_paragraph(correspondencia_saliente.cite)
        doc.add_paragraph(correspondencia_saliente.remitente.nombre)
        doc.add_paragraph(correspondencia_saliente.remitente.apellido)
        doc.add_paragraph(correspondencia_saliente.remitente.cargo)
        doc.add_paragraph(str(correspondencia_saliente.remitente.institucion))    
        doc.add_paragraph(f"Ref.: {correspondencia_saliente.referencia}")
        doc.add_paragraph(correspondencia_saliente.descripcion)


        # Guardar el archivo en un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Devolver el archivo como respuesta
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename=correspondencia_{correspondencia_saliente.cite}.docx'
        return response

    generar_documento_word.short_description = "Generar documento Word"
@admin.register(DescargaDocumento)
class DescargaDocumentoAdmin(admin.ModelAdmin):
    list_display = ('usuario', 'documento', 'fecha_descarga')
    list_filter = ('usuario', 'documento')
    search_fields = ('usuario__nombre', 'documento__cite')
    readonly_fields = ('fecha_descarga',)

  



# Registra los demás modelos sin personalización

admin.site.register(TipoDocumento)

