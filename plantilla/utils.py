from io import BytesIO
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from django.utils.timezone import now
import os

def generar_documento_word(correspondencia_saliente):
    """Genera un documento Word a partir de un objeto CorrespondenciaSaliente."""

    doc = Document()
    doc.add_paragraph(f"La Paz {correspondencia_saliente.fecha_envio.strftime('%Y-%m-%d')}")
    doc.add_paragraph(correspondencia_saliente.cite)
    doc.add_paragraph(correspondencia_saliente.remitente.nombre)
    doc.add_paragraph(correspondencia_saliente.remitente.apellido)
    doc.add_paragraph(correspondencia_saliente.remitente.cargo)
    doc.add_paragraph(str(correspondencia_saliente.remitente.institucion))
    doc.add_paragraph("De nuestra mayor consideracion:")
    doc.add_paragraph(f"Ref.: {correspondencia_saliente.referencia}")
    doc.add_paragraph(correspondencia_saliente.descripcion)

    # Guardar el archivo en un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Devolver el archivo como respuesta
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=correspondencia_{correspondencia_saliente.cite}.docx'
    
    return response

def generar_documento_interno(tipo, numero, gestion, contenido, destinatario):
    """
    Genera un archivo Word con un formato específico según el tipo de documento.
    
    :param tipo: Tipo de documento (Ej: Comunicado, Convocatoria, etc.)
    :param numero: Número secuencial.
    :param gestion: Año del documento.
    :param contenido: Cuerpo del documento.
    :param destinatario: Persona o entidad destinataria.
    :return: Ruta del archivo generado.
    """
    
    # Crear el documento Word
    doc = Document()
    
    # 📌 Agregar título según el tipo de documento
    doc.add_heading(tipo.upper(), level=1)
    
    # 📌 Fecha del documento
    fecha = now().strftime("%d de %B de %Y")
    doc.add_paragraph(f"La Paz, {fecha}")

    # 📌 Número del documento (CITE)
    cite = f"{tipo[:3].upper()}-{numero:03d}/{gestion}"
    doc.add_paragraph(f"CITE: {cite}")

    # 📌 Espacio antes del contenido
    doc.add_paragraph("\n")

    # 📌 Destinatario
    doc.add_paragraph(f"Para: {destinatario}")

    # 📌 Espacio antes del contenido
    doc.add_paragraph("\n")

    # 🔹 PERSONALIZAR FORMATO SEGÚN EL TIPO DE DOCUMENTO 🔹
    if tipo.lower() == "comunicado":
        doc.add_paragraph("De nuestra consideración:")
        doc.add_paragraph("\n")
        p = doc.add_paragraph(contenido)
        p.runs[0].bold = True  # Resaltar el comunicado en negrita

    elif tipo.lower() == "convocatoria":
        doc.add_paragraph("Se convoca a:")
        doc.add_paragraph("\n")
        p = doc.add_paragraph(contenido)
        p.runs[0].italic = True  # Texto en cursiva para énfasis

    elif tipo.lower() == "aviso":
        doc.add_paragraph("Atención:")
        doc.add_paragraph("\n")
        p = doc.add_paragraph(contenido)
        p.runs[0].font.size = Pt(14)  # Tamaño de letra más grande para avisos

    elif tipo.lower() == "informe":
        doc.add_paragraph("Asunto:")
        doc.add_paragraph("\n")
        p = doc.add_paragraph(contenido)
        p.runs[0].font.size = Pt(12)
        p.runs[0].underline = True  # Subrayado en informes

    else:
        doc.add_paragraph(contenido)  # Formato estándar si no es un tipo conocido

    # 📌 Espacio antes de la firma
    doc.add_paragraph("\n\nAtentamente,\n")

    # 📌 Guardar el archivo en la carpeta `media/documentos/`
    ruta_carpeta = "media/documentos/"
    os.makedirs(ruta_carpeta, exist_ok=True)

    nombre_archivo = f"{tipo.upper()}_{numero:03d}_{gestion}.docx"
    ruta_archivo = os.path.join(ruta_carpeta, nombre_archivo)
    
    doc.save(ruta_archivo)
    
    return ruta_archivo